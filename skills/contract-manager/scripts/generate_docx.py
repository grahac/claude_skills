#!/usr/bin/env python3
"""
Convert a JSON contract definition into a professionally formatted .docx file.

Uses Calibri, Word heading styles for navigation, Word auto-numbering for
lettered lists (a)(b)(c), and clean signature blocks.

Usage:
    python generate_docx.py --input contract.json --output contract.docx
"""

import argparse
import json
import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print(
        "Error: python-docx is not installed.\n"
        "Install it with: pip install python-docx"
    )
    sys.exit(1)


REQUIRED_FIELDS = ["title", "date", "sections"]
FONT_NAME = "Calibri"


def validate_contract(data):
    """Validate that all required top-level fields exist in the contract JSON."""
    missing = [field for field in REQUIRED_FIELDS if field not in data]
    if missing:
        print(f"Error: Missing required fields in JSON: {', '.join(missing)}")
        sys.exit(1)
    for i, party in enumerate(data.get("parties", [])):
        for key in ("name", "role", "address"):
            if key not in party:
                print(f"Error: Party at index {i} is missing required field '{key}'.")
                sys.exit(1)
    for i, section in enumerate(data["sections"]):
        for key in ("number", "heading", "content"):
            if key not in section:
                print(f"Error: Section at index {i} is missing required field '{key}'.")
                sys.exit(1)
    for i, sig in enumerate(data.get("signature_blocks", [])):
        for key in ("party", "name_line", "title_line"):
            if key not in sig:
                print(f"Error: Signature block at index {i} missing '{key}'.")
                sys.exit(1)


# ---------------------------------------------------------------------------
# Font / style helpers
# ---------------------------------------------------------------------------

def set_run_font(run, size_pt=11, bold=False, italic=False):
    """Apply font settings to a run."""
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)


def ensure_style_font(style):
    """Set font on a style element for all font slots."""
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), FONT_NAME)


def set_paragraph_spacing(paragraph, space_before_pt=0, space_after_pt=6):
    """Set spacing on a paragraph."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    pf.line_spacing = 1.15


def configure_styles(doc):
    """Set up document styles — Normal, Heading 1, Heading 2, Signature."""
    # Normal
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    ensure_style_font(normal)

    # Heading 1 — major contract sections
    h1 = doc.styles["Heading 1"]
    h1.font.name = FONT_NAME
    h1.font.size = Pt(12)
    h1.font.bold = True
    h1.font.color.rgb = None  # automatic/black
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(2)
    h1.paragraph_format.line_spacing = 1.15
    h1.paragraph_format.keep_with_next = True
    ensure_style_font(h1)

    # Heading 2 — subsections
    h2 = doc.styles["Heading 2"]
    h2.font.name = FONT_NAME
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = None
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(2)
    h2.paragraph_format.line_spacing = 1.15
    h2.paragraph_format.left_indent = Inches(0.5)
    h2.paragraph_format.keep_with_next = True
    ensure_style_font(h2)

    # Signature — single-spaced, no extra spacing
    try:
        sig = doc.styles.add_style("Signature", WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        sig = doc.styles["Signature"]
    sig.base_style = normal
    sig.font.name = FONT_NAME
    sig.font.size = Pt(11)
    sig.paragraph_format.space_before = Pt(0)
    sig.paragraph_format.space_after = Pt(0)
    sig.paragraph_format.line_spacing = 1.0
    ensure_style_font(sig)


# ---------------------------------------------------------------------------
# Word auto-numbering for (a)(b)(c) lists
# ---------------------------------------------------------------------------

def create_letter_abstract_num(abstract_id, left_twips, hanging_twips):
    """Build an abstractNum XML element for (a)(b)(c) numbering."""
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_id))

    mt = OxmlElement("w:multiLevelType")
    mt.set(qn("w:val"), "singleLevel")
    abstract_num.append(mt)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "lowerLetter")
    lvl.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "(%1)")
    lvl.append(lvl_text)

    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    pPr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left_twips))
    ind.set(qn("w:hanging"), str(hanging_twips))
    pPr.append(ind)
    lvl.append(pPr)

    abstract_num.append(lvl)
    return abstract_num


def setup_letter_numbering(doc):
    """Create (a)(b)(c) numbering definitions for two indent levels.

    Returns (top_abstract_id, sub_abstract_id) for top-level and subsection lists.
    """
    # Ensure numbering part exists by adding/removing a dummy list paragraph
    temp = doc.add_paragraph(style="List Bullet")
    doc.element.body.remove(temp._element)

    numbering = doc.part.numbering_part.element

    existing = numbering.findall(qn("w:abstractNum"))
    max_id = max(
        (int(e.get(qn("w:abstractNumId"))) for e in existing), default=-1
    )

    # Top-level lists: label at 0.25", text at 0.5"
    top_id = max_id + 1
    top_abs = create_letter_abstract_num(top_id, left_twips=720, hanging_twips=360)

    # Subsection lists: label at 0.75", text at 1.0"
    sub_id = max_id + 2
    sub_abs = create_letter_abstract_num(sub_id, left_twips=1440, hanging_twips=360)

    # Insert abstractNums before any num elements
    first_num = numbering.find(qn("w:num"))
    if first_num is not None:
        idx = list(numbering).index(first_num)
        numbering.insert(idx, sub_abs)
        numbering.insert(idx, top_abs)
    else:
        numbering.append(top_abs)
        numbering.append(sub_abs)

    return top_id, sub_id


def get_new_num_id(doc, abstract_id):
    """Create a new num instance that restarts numbering from (a)."""
    numbering = doc.part.numbering_part.element

    existing_nums = numbering.findall(qn("w:num"))
    max_num_id = max(
        (int(e.get(qn("w:numId"))) for e in existing_nums), default=0
    )
    num_id = max_num_id + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))

    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)

    # Restart override — always start from (a)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)

    numbering.append(num)
    return num_id


def apply_list_numbering(paragraph, num_id):
    """Apply a numbering definition to a paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId_elm = OxmlElement("w:numId")
    numId_elm.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId_elm)
    pPr.append(numPr)


# ---------------------------------------------------------------------------
# Page numbers
# ---------------------------------------------------------------------------

def add_page_numbers(doc):
    """Add centered page numbers to the document footer."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    set_run_font(run, 10)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._element.append(begin)

    run2 = paragraph.add_run()
    set_run_font(run2, 10)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._element.append(instr)

    run3 = paragraph.add_run()
    set_run_font(run3, 10)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run3._element.append(end)


# ---------------------------------------------------------------------------
# Document building — title, parties, recitals
# ---------------------------------------------------------------------------

def add_title(doc, title_text, date_text):
    """Add the contract title and date."""
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title_para, space_after_pt=2)
    set_run_font(title_para.add_run(title_text), 14, bold=True)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(date_para, space_after_pt=12)
    set_run_font(date_para.add_run(date_text), 11)


def add_parties(doc, parties):
    """Add the BETWEEN: parties section."""
    header_para = doc.add_paragraph()
    set_paragraph_spacing(header_para, space_before_pt=6, space_after_pt=10)
    set_run_font(header_para.add_run("BETWEEN:"), 11, bold=True)

    for i, party in enumerate(parties):
        para = doc.add_paragraph()
        set_paragraph_spacing(para, space_before_pt=8, space_after_pt=12)
        para.paragraph_format.left_indent = Inches(0.5)

        set_run_font(para.add_run(party["name"]), 11, bold=True)
        set_run_font(para.add_run(f' (the "{party["role"]}")'), 11)
        set_run_font(para.add_run(f"\n{party['address']}"), 11)

        if i < len(parties) - 1:
            and_para = doc.add_paragraph()
            and_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(and_para, space_before_pt=6, space_after_pt=6)
            set_run_font(and_para.add_run("and"), 11)

    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, space_after_pt=12)


def add_recitals(doc, recitals_text):
    """Add the recitals section in italic."""
    if not recitals_text:
        return
    para = doc.add_paragraph()
    set_paragraph_spacing(para)
    set_run_font(para.add_run(recitals_text), 11, italic=True)

    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, space_after_pt=6)


# ---------------------------------------------------------------------------
# Content parsing — detects (a), (b), (c) items
# ---------------------------------------------------------------------------

def parse_content_blocks(content):
    """Split section content into prose blocks and lettered list items.

    Returns list of ('prose', text) or ('item', letter, text) tuples.
    """
    blocks = []
    paragraphs = re.split(r"\n\n+", content.strip())

    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        # Paragraph starts with "(a) ..." — may contain \n(b), \n(c) too
        if re.match(r"^\([a-z]\)\s", para_text):
            parts = re.split(r"\n\(([a-z])\)\s", para_text)
            # First part is "(a) text..."
            blocks.append(("item", parts[0][1], parts[0][4:].strip()))
            for j in range(1, len(parts), 2):
                letter = parts[j]
                text = parts[j + 1].strip() if j + 1 < len(parts) else ""
                blocks.append(("item", letter, text))
        else:
            # Check for items split across lines: "intro:\n(a) ...\n(b) ..."
            inline_split = re.split(r"\n\(([a-z])\)\s", para_text)
            if len(inline_split) > 1:
                intro = inline_split[0].strip()
                if intro:
                    blocks.append(("prose", intro))
                for j in range(1, len(inline_split), 2):
                    letter = inline_split[j]
                    text = (
                        inline_split[j + 1].strip()
                        if j + 1 < len(inline_split)
                        else ""
                    )
                    blocks.append(("item", letter, text))
            else:
                blocks.append(("prose", para_text))

    return blocks


# ---------------------------------------------------------------------------
# Contract sections — headings + content with auto-numbered lists
# ---------------------------------------------------------------------------

def add_content_blocks(doc, content, base_indent, abstract_id):
    """Render content as prose paragraphs and Word auto-numbered lists."""
    blocks = parse_content_blocks(content)

    # If there are list items, create a new num instance (restarts from 'a')
    has_items = any(b[0] == "item" for b in blocks)
    num_id = get_new_num_id(doc, abstract_id) if has_items else None

    for block in blocks:
        if block[0] == "prose":
            para = doc.add_paragraph()
            set_paragraph_spacing(para, space_before_pt=0, space_after_pt=6)
            para.paragraph_format.left_indent = base_indent
            set_run_font(para.add_run(block[1]), 11)
        elif block[0] == "item":
            para = doc.add_paragraph()
            set_paragraph_spacing(para, space_before_pt=0, space_after_pt=3)
            apply_list_numbering(para, num_id)
            # Word auto-generates the "(a)" label; just add the text
            set_run_font(para.add_run(block[2]), 11)


def add_sections(doc, sections, top_abstract_id, sub_abstract_id):
    """Add numbered contract sections using Heading 1/2 styles."""
    for section in sections:
        number = section["number"]
        heading_text = f"{number}. {section['heading']}"
        is_subsection = "." in number
        is_exhibit = number.startswith("EXHIBIT") or re.match(r"^[A-Z]\.", number)

        # Choose heading style and indent level
        if is_subsection and not is_exhibit:
            style = "Heading 2"
            base_indent = Inches(0.5)
            abstract_id = sub_abstract_id
        elif is_exhibit and is_subsection:
            style = "Heading 2"
            base_indent = Inches(0)
            abstract_id = top_abstract_id
        else:
            style = "Heading 1"
            base_indent = Inches(0)
            abstract_id = top_abstract_id

        heading_para = doc.add_paragraph(heading_text, style=style)

        # Exhibit subsections: remove the default Heading 2 indent
        if is_exhibit and is_subsection:
            heading_para.paragraph_format.left_indent = Inches(0)

        # Skip empty content (section headers with no body text)
        if not section["content"]:
            continue

        add_content_blocks(doc, section["content"], base_indent, abstract_id)


# ---------------------------------------------------------------------------
# Signature blocks — single-spaced, no labels, date right-aligned
# ---------------------------------------------------------------------------

def add_right_tab(paragraph, pos_inches=6.5):
    """Add a right-aligned tab stop to a paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(int(pos_inches * 1440)))
    tab.set(qn("w:leader"), "none")
    tabs.append(tab)
    pPr.append(tabs)


def add_signature_blocks(doc, signature_blocks):
    """Add signature blocks — single-spaced, no labels, date right-aligned."""
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, space_after_pt=12)
    set_run_font(
        spacer.add_run(
            "IN WITNESS WHEREOF, the parties have executed this Agreement "
            "as of the date first written above."
        ),
        11,
    )

    for sig in signature_blocks:
        # Space before each block
        sp = doc.add_paragraph(style="Signature")
        sp.paragraph_format.space_after = Pt(30)

        # Party name (bold)
        party_para = doc.add_paragraph(style="Signature")
        set_run_font(party_para.add_run(sig["party"]), 11, bold=True)

        # Signature line (with space above for actual signing)
        sig_para = doc.add_paragraph(style="Signature")
        sig_para.paragraph_format.space_before = Pt(24)
        set_run_font(sig_para.add_run("_" * 40), 11)

        # Name
        name_para = doc.add_paragraph(style="Signature")
        set_run_font(name_para.add_run(sig["name_line"]), 11)

        # Title + Date on same line (date right-aligned)
        td_para = doc.add_paragraph(style="Signature")
        add_right_tab(td_para)
        set_run_font(td_para.add_run(sig["title_line"]), 11)
        set_run_font(td_para.add_run("\t"), 11)
        set_run_font(td_para.add_run("Date: _______________"), 11)


# ---------------------------------------------------------------------------
# Main generation
# ---------------------------------------------------------------------------

def generate_docx(data, output_path):
    """Generate the .docx document from validated contract data."""
    doc = Document()

    # Set 1-inch margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    configure_styles(doc)
    top_abs_id, sub_abs_id = setup_letter_numbering(doc)
    add_page_numbers(doc)
    add_title(doc, data["title"], data["date"])
    if data.get("parties"):
        add_parties(doc, data["parties"])
    add_recitals(doc, data.get("recitals", ""))
    add_sections(doc, data["sections"], top_abs_id, sub_abs_id)
    if data.get("signature_blocks"):
        add_signature_blocks(doc, data["signature_blocks"])

    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(
        description="Convert a JSON contract definition into a formatted .docx file."
    )
    parser.add_argument("--input", required=True, help="Path to input JSON file.")
    parser.add_argument("--output", required=True, help="Path for output .docx file.")
    args = parser.parse_args()

    try:
        with open(args.input, "r", encoding="utf-8") as f:
            data = json.load(f)
    except FileNotFoundError:
        print(f"Error: Input file not found: {args.input}")
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"Error: Invalid JSON: {e}")
        sys.exit(1)

    validate_contract(data)

    output_dir = os.path.dirname(args.output)
    if output_dir and not os.path.isdir(output_dir):
        print(f"Error: Output directory does not exist: {output_dir}")
        sys.exit(1)

    generate_docx(data, args.output)
    print(f"Contract document generated: {args.output}")


if __name__ == "__main__":
    main()
